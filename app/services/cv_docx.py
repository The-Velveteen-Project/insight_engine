"""
Render a TailoredCV as a .docx (Phase 2.6).

Deterministic layout, no template dependency: identity block, headline,
summary, then the same sections as the Markdown render. ATS-friendly:
plain headings, bullet paragraphs, one column, no tables or text boxes.
"""

from __future__ import annotations

from io import BytesIO

from docx import Document as new_document
from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.schemas.cv import CVEntry, TailoredCV


def _set_base_font(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)


def _heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(2)


def _bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(text.strip(), style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(1)


def _entry(document: Document, entry: CVEntry) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(entry.title.strip())
    run.bold = True
    if entry.subtitle:
        paragraph.add_run(f" · {entry.subtitle.strip()}")
    paragraph.paragraph_format.space_after = Pt(1)
    for bullet in entry.bullets:
        _bullet(document, bullet)


def render_docx(cv: TailoredCV, *, identity_block: str) -> bytes:
    document = new_document()
    _set_base_font(document)
    for section in document.sections:
        section.top_margin = Pt(40)
        section.bottom_margin = Pt(40)
        section.left_margin = Pt(50)
        section.right_margin = Pt(50)

    identity_lines = [line for line in identity_block.splitlines() if line.strip()]
    if identity_lines:
        name = document.add_paragraph()
        name_run = name.add_run(identity_lines[0].strip())
        name_run.bold = True
        name_run.font.size = Pt(16)
        name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name.paragraph_format.space_after = Pt(0)
        for line in identity_lines[1:]:
            contact = document.add_paragraph(line.strip())
            contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact.paragraph_format.space_after = Pt(0)

    headline = document.add_paragraph()
    headline_run = headline.add_run(cv.headline.strip())
    headline_run.bold = True
    headline_run.font.size = Pt(12)
    headline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    headline.paragraph_format.space_before = Pt(6)

    document.add_paragraph(cv.summary.strip())

    _heading(document, "Selected projects")
    for entry in cv.highlighted_projects:
        _entry(document, entry)
    _heading(document, "Experience")
    for entry in cv.experience:
        _entry(document, entry)
    _heading(document, "Education")
    for item in cv.education:
        _bullet(document, item)
    if cv.publications:
        _heading(document, "Publications and manuscripts")
        for item in cv.publications:
            _bullet(document, item)
    if cv.distinctions:
        _heading(document, "Fellowships and awards")
        for item in cv.distinctions:
            _bullet(document, item)
    _heading(document, "Skills")
    for item in cv.skills:
        _bullet(document, item)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
